"""Multi-format document parser (PDF, DOCX, XLSX, image OCR, text, email) with semantic chunking."""

import email
import re
from pathlib import Path

import fitz  # PyMuPDF
import pytesseract
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image


class DocumentParser:
    """Multi-format document parser with semantic chunking.

    Supports: PDF, DOCX, XLSX, images (OCR), plain text, .eml email files.
    Chunks documents by semantic boundaries (headings, sections, tables)
    rather than fixed token windows, preserving context for requirement extraction.
    """

    MAX_CHUNK_CHARS = 3000
    MIN_CHUNK_CHARS = 50

    def parse(self, file_path: str, file_type: str) -> list[dict]:
        """Parse a document into semantic chunks.

        Args:
            file_path: Absolute path to the document file
            file_type: One of: pdf, docx, xlsx, image, text, email

        Returns:
            List of chunk dicts with keys: id, content, page_number,
            section_title, heading_hierarchy, chunk_type, source_filename
        """
        # Validate the file_type FIRST so callers can detect typos without
        # needing the file on disk. Disk-existence is the next check.
        parser_map = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "xlsx": self._parse_xlsx,
            "image": self._parse_image,
            "text": self._parse_text,
            "email": self._parse_email,
        }
        parser_fn = parser_map.get(file_type)
        if not parser_fn:
            raise ValueError(f"Unsupported file type: {file_type}")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        raw_chunks = parser_fn(path)

        # Post-process: assign IDs, filter empty, enforce size limits
        processed = []
        for idx, chunk in enumerate(raw_chunks):
            content = chunk.get("content", "").strip()
            if len(content) < self.MIN_CHUNK_CHARS:
                continue

            # Split oversized chunks
            if len(content) > self.MAX_CHUNK_CHARS:
                sub_chunks = self._split_large_chunk(content)
                for sub_idx, sub_content in enumerate(sub_chunks):
                    processed.append({
                        "id": f"{path.stem}-{idx:04d}-{sub_idx:02d}",
                        "content": sub_content,
                        "page_number": chunk.get("page_number"),
                        "section_title": chunk.get("section_title"),
                        "heading_hierarchy": chunk.get("heading_hierarchy", []),
                        "chunk_type": chunk.get("chunk_type", "paragraph"),
                        "source_filename": path.name,
                        "section_info": chunk.get("section_title", f"Page {chunk.get('page_number', '?')}"),
                    })
            else:
                processed.append({
                    "id": f"{path.stem}-{idx:04d}",
                    "content": content,
                    "page_number": chunk.get("page_number"),
                    "section_title": chunk.get("section_title"),
                    "heading_hierarchy": chunk.get("heading_hierarchy", []),
                    "chunk_type": chunk.get("chunk_type", "paragraph"),
                    "source_filename": path.name,
                    "section_info": chunk.get("section_title", f"Page {chunk.get('page_number', '?')}"),
                })

        return processed

    def _parse_pdf(self, path: Path) -> list[dict]:
        """Parse PDF using PyMuPDF with heading-aware chunking."""
        doc = fitz.open(str(path))
        try:
            chunks: list[dict] = []
            current_section = ""
            current_content = []
            current_page = 1

            for page_num in range(len(doc)):
                page = doc[page_num]
                blocks = page.get_text("dict")["blocks"]

                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        for line in block.get("lines", []):
                            text = "".join(span["text"] for span in line.get("spans", []))
                            font_size = max(
                                (span.get("size", 12) for span in line.get("spans", [])),
                                default=12,
                            )

                            # Detect headings by font size
                            if font_size > 14 and text.strip():
                                # Save previous section
                                if current_content:
                                    chunks.append({
                                        "content": "\n".join(current_content),
                                        "page_number": current_page,
                                        "section_title": current_section,
                                        "chunk_type": "paragraph",
                                    })
                                current_section = text.strip()
                                current_content = []
                                current_page = page_num + 1
                            elif text.strip():
                                current_content.append(text.strip())

            # Final section
            if current_content:
                chunks.append({
                    "content": "\n".join(current_content),
                    "page_number": current_page,
                    "section_title": current_section,
                    "chunk_type": "paragraph",
                })

            return chunks
        finally:
            doc.close()

    def _parse_docx(self, path: Path) -> list[dict]:
        """Parse DOCX with heading hierarchy tracking."""
        doc = DocxDocument(str(path))
        chunks: list[dict] = []
        current_headings: list[str] = []
        current_content: list[str] = []
        current_section = ""

        for para in doc.paragraphs:
            style = para.style.name if para.style else ""

            if style.startswith("Heading"):
                # Save previous chunk
                if current_content:
                    chunks.append({
                        "content": "\n".join(current_content),
                        "section_title": current_section,
                        "heading_hierarchy": list(current_headings),
                        "chunk_type": "paragraph",
                    })
                    current_content = []

                # Update heading hierarchy
                level = int(style.replace("Heading ", "")) if style != "Heading" else 1
                current_headings = current_headings[:level - 1]
                current_headings.append(para.text.strip())
                current_section = para.text.strip()

            elif para.text.strip():
                current_content.append(para.text.strip())

        # Final chunk
        if current_content:
            chunks.append({
                "content": "\n".join(current_content),
                "section_title": current_section,
                "heading_hierarchy": list(current_headings),
                "chunk_type": "paragraph",
            })

        # Parse tables
        for idx, table in enumerate(doc.tables):
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            if rows_text:
                chunks.append({
                    "content": "\n".join(rows_text),
                    "section_title": f"Table {idx + 1}",
                    "chunk_type": "table",
                })

        return chunks

    def _parse_xlsx(self, path: Path) -> list[dict]:
        """Parse Excel workbook — each sheet becomes chunks based on data regions."""
        wb = load_workbook(str(path), read_only=True, data_only=True)
        try:
            chunks: list[dict] = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows_text: list[str] = []

                for row in ws.iter_rows(values_only=True):
                    cells = [str(cell) if cell is not None else "" for cell in row]
                    row_text = " | ".join(c for c in cells if c)
                    if row_text.strip():
                        rows_text.append(row_text)

                if rows_text:
                    # Chunk by groups of rows
                    for i in range(0, len(rows_text), 20):
                        batch = rows_text[i : i + 20]
                        chunks.append({
                            "content": "\n".join(batch),
                            "section_title": f"Sheet: {sheet_name}",
                            "chunk_type": "table",
                        })

            return chunks
        finally:
            wb.close()

    def _parse_image(self, path: Path) -> list[dict]:
        """Parse image via OCR (pytesseract)."""
        with Image.open(str(path)) as image:
            if image.format and image.format.upper() not in {"PNG", "JPEG", "TIFF", "BMP", "GIF"}:
                raise ValueError(f"Unsupported image format: {image.format}")
            text = pytesseract.image_to_string(image)

        if text.strip():
            return [{
                "content": text.strip(),
                "section_title": f"OCR: {path.name}",
                "chunk_type": "image_description",
            }]
        return []

    def _parse_text(self, path: Path) -> list[dict]:
        """Parse plain text file with paragraph-based chunking."""
        content = path.read_text(encoding="utf-8", errors="ignore")
        paragraphs = re.split(r"\n{2,}", content)
        chunks: list[dict] = []

        for idx, para in enumerate(paragraphs):
            if para.strip():
                chunks.append({
                    "content": para.strip(),
                    "section_title": f"Section {idx + 1}",
                    "chunk_type": "paragraph",
                })

        return chunks

    def _parse_email(self, path: Path) -> list[dict]:
        """Parse .eml email file — extract subject, body, and attachments metadata."""
        with open(str(path), "r", encoding="utf-8", errors="ignore") as f:
            msg = email.message_from_file(f)

        chunks: list[dict] = []
        subject = msg.get("Subject", "No Subject")
        sender = msg.get("From", "Unknown")
        date = msg.get("Date", "Unknown")

        # Email header as context
        header = f"From: {sender}\nDate: {date}\nSubject: {subject}"

        # Extract body
        body_parts: list[str] = []
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        body_parts.append(payload.decode("utf-8", errors="ignore"))
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                body_parts.append(payload.decode("utf-8", errors="ignore"))

        full_body = "\n".join(body_parts)
        if full_body.strip():
            chunks.append({
                "content": f"{header}\n\n{full_body.strip()}",
                "section_title": f"Email: {subject}",
                "chunk_type": "paragraph",
            })

        return chunks

    def _split_large_chunk(self, content: str) -> list[str]:
        """Split oversized content at paragraph or sentence boundaries."""
        paragraphs = content.split("\n")
        sub_chunks: list[str] = []
        current: list[str] = []
        current_len = 0

        for para in paragraphs:
            if current_len + len(para) > self.MAX_CHUNK_CHARS and current:
                sub_chunks.append("\n".join(current))
                current = []
                current_len = 0
            current.append(para)
            current_len += len(para)

        if current:
            sub_chunks.append("\n".join(current))

        return sub_chunks
