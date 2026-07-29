"""
Requirements Document Extractor
Extracts text from .docx and .xlsx files and outputs structured JSON.
Replaces in-context document reading to reduce token consumption.

Usage:
  python rom_extract.py --path "C:\\path\\to\\file.docx" [--format json|text]
  python rom_extract.py --dir "C:\\path\\to\\RequirementDocs" --list

Output (--list mode): JSON list of documents in directory
Output (--path mode): JSON with extracted content
"""
import json
import sys
import argparse
from pathlib import Path


def list_documents(directory: str) -> list:
    """List requirements documents in a directory."""
    dir_path = Path(directory)
    if not dir_path.exists():
        return []
    extensions = {".docx", ".xlsx", ".pdf", ".txt", ".md"}
    docs = []
    for i, f in enumerate(sorted(dir_path.iterdir()), 1):
        if f.is_file() and f.suffix.lower() in extensions:
            docs.append({
                "number": i,
                "name": f.name,
                "path": str(f),
                "extension": f.suffix.lower(),
                "size_kb": round(f.stat().st_size / 1024, 1),
            })
    return docs


def extract_docx(file_path: str) -> dict:
    """Extract text and tables from a .docx file."""
    try:
        import docx
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    doc = docx.Document(file_path)
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append({
                "text": text,
                "style": p.style.name if p.style else "",
            })

    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(cells)
        tables.append({"table_index": ti, "rows": rows})

    # Build a concise summary (first 500 chars of body text)
    full_text = "\n".join(p["text"] for p in paragraphs)
    summary = full_text[:500] + ("..." if len(full_text) > 500 else "")

    return {
        "file": Path(file_path).name,
        "type": "docx",
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "summary": summary,
        "paragraphs": paragraphs,
        "tables": tables,
    }


def extract_xlsx(file_path: str) -> dict:
    """Extract data from an .xlsx file."""
    try:
        import openpyxl
    except ImportError:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    wb = openpyxl.load_workbook(file_path, data_only=True)
    sheets = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(c is not None and str(c).strip() != "" for c in row):
                rows.append(["" if c is None else str(c) for c in row])
        sheets.append({
            "sheet_name": name,
            "row_count": len(rows),
            "rows": rows,
        })

    return {
        "file": Path(file_path).name,
        "type": "xlsx",
        "sheet_count": len(sheets),
        "sheets": sheets,
    }


def extract_text(file_path: str) -> dict:
    """Extract content from a plain text or markdown file."""
    path = Path(file_path)
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    return {
        "file": path.name,
        "type": path.suffix.lower().lstrip("."),
        "char_count": len(content),
        "content": content,
    }


def extract_file(file_path: str) -> dict:
    """Route extraction based on file type."""
    path = Path(file_path)
    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_docx(file_path)
    elif ext == ".xlsx":
        return extract_xlsx(file_path)
    elif ext in (".txt", ".md"):
        return extract_text(file_path)
    else:
        return {"error": f"Unsupported file type: {ext}"}


def main():
    parser = argparse.ArgumentParser(description="Requirements Document Extractor")
    parser.add_argument("--path", help="Path to a single document to extract")
    parser.add_argument("--dir", help="Path to requirements directory")
    parser.add_argument("--list", action="store_true", help="List documents in directory")
    parser.add_argument("--format", choices=["json", "text"], default="json", help="Output format")
    parser.add_argument("--select", help="Comma-separated document numbers or 'all'")
    args = parser.parse_args()

    sys.stdout.reconfigure(encoding="utf-8")

    if args.dir and args.list:
        docs = list_documents(args.dir)
        print(json.dumps({"directory": args.dir, "document_count": len(docs), "documents": docs}, indent=2))
        return

    if args.dir and args.select:
        docs = list_documents(args.dir)
        if args.select.lower() == "all":
            selected = docs
        else:
            indices = [int(x.strip()) for x in args.select.split(",")]
            selected = [d for d in docs if d["number"] in indices]

        results = []
        for doc in selected:
            result = extract_file(doc["path"])
            results.append(result)
        print(json.dumps({"selected_count": len(results), "extractions": results}, indent=2))
        return

    if args.path:
        result = extract_file(args.path)
        if args.format == "text" and "paragraphs" in result:
            # Output plain text for simpler consumption
            for p in result["paragraphs"]:
                print(p["text"])
            for t in result.get("tables", []):
                print(f"\n--- Table {t['table_index']} ---")
                for row in t["rows"]:
                    print(" | ".join(row))
        else:
            print(json.dumps(result, indent=2))
        return

    parser.print_help()


if __name__ == "__main__":
    main()
